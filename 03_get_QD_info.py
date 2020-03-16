# -*- coding: utf-8 -*-
"""
Created on Fri Sep  6 14:23:28 2019

@author: gbm294
Use existing csv file with Report Spec doc paths to
Get IRB and PI info from Report spec Docs (on U: drive)
"""

import sys, os
import tkinter
from tkinter import filedialog #GUI

import pandas as pd
pd.set_option("display.max_colwidth", 10000)  ##prevents strings from being truncated.
from docx import Document

import read_report_spec_v2



def get_file_list(what_file,parent_dir):
    #get input file
    root = tkinter.Tk()  #access windowing system
    effort_file = filedialog.askopenfilename(initialdir = parent_dir,title=what_file)
    user_dir = os.path.dirname(effort_file)
    #os.chdir(user_dir)  #specify directory
    user_dir_name = os.path.basename(os.path.dirname(effort_file))
    file_path = effort_file #[0]
    root.destroy() #kill Tkinter
    
    if len(effort_file) <1 :
        print("you probably didn't choose any files")
        root.destroy() #kill Tkinter
        sys.exit()
    
    return file_path , user_dir


def process_csv(in_path):
    df = pd.read_csv(in_path, sep=',', escapechar='\\', quotechar='"', encoding = "ISO-8859-1", keep_default_na =False)  #encoding = "ISO-8859-1"
    return df

def process_word_doc(in_path):
    print(in_path)
    wordDoc = Document(in_path)
#    cells = []
#    for table in wordDoc.tables:
#        for row in table.rows:
#            for cell in row.cells:
#                cells.append(cell.text)
    study_info = {}
    irb = ''
    title = ''
    p_i = ''
    dept = ''
    for p in wordDoc.paragraphs:
        pt = p.text
        if 'IRB Submission ID number:' in pt:
            irb = pt.split(':',1)[1].strip()
            study_info['IRB Submission ID number'] = irb
            continue
        if 'Study Title:' in pt:
            title = pt.split(':',1)[1].strip()
            study_info['Study Title'] = title
            continue
        if 'Principal Investigator (PI):' in pt:
            p_i = pt.split(':',1)[1].strip()
            study_info['Principal Investigator'] = p_i
            continue
        if 'PI Department:' in pt:
            dept = pt.split(':',1)[1].strip()
            study_info['PI Department'] = dept
            break
        


    return study_info


def main():
    print("Let's go!")
    
    ## path for new RITM list
#    new_ritm_list_path = 'U:/UWHealth/EA/SpecialShares/DM/CRDS/AdHocQueries/ICTR Metrics Dashboard/Effort_tracking_REDCap/2019_12_17/new_ritms_2019-12-17_python.csv'
#    qd_path_list = 'U:/UWHealth/EA/SpecialShares/DM/CRDS/AdHocQueries/ICTR Metrics Dashboard/Effort_tracking_REDCap/2019_12_17/QD_path_list.csv'
#    qd_details = 'U:/UWHealth/EA/SpecialShares/DM/CRDS/AdHocQueries/ICTR Metrics Dashboard/Effort_tracking_REDCap/2019_12_17/QD_details.csv'
    
    new_ritm_list_path, new_ritm_list_dir = get_file_list('New RITMs List' ,'U:/UWHealth/EA/SpecialShares/DM/CRDS/AdHocQueries/ICTR Metrics Dashboard/Manual_effort_calcs_gm')
    qd_path_list, qd_path_list_dir = get_file_list('QD Path List' ,'U:/UWHealth/EA/SpecialShares/DM/CRDS/AdHocQueries/ICTR Metrics Dashboard/Manual_effort_calcs_gm')
    
    os.chdir(new_ritm_list_dir)
    os.chdir("..")
    working_dir = '%s/03_get_QD_info_Output' % os.getcwd()
    os.chdir(working_dir)
    qd_details = 'QD_details.csv'


    ##get list of new RITMs in df
    new_ritm_list_path_df = process_csv(new_ritm_list_path)
    qd_path_list_df = process_csv(qd_path_list)

    
    out_info = []
    

    ## Work through csv file and pull in IRB and PI data for the RITM
    for index, row in new_ritm_list_path_df.iterrows():
        try:
            #print(row['SN_RITM'], row['short_description'])
            ## open matching qd and get IRB and PI info from word doc
            
            #merge_df = pd.merge(sn_df, rc_df, left_on='number', right_on='RITM', how='left')
            merge_df = qd_path_list_df[qd_path_list_df['ritm'] == row['SN_RITM']].head(1)
            L = merge_df.shape[0]
            if L <1:    
                print('no folder')
                out_info.append({'ritm':row['SN_RITM']
                                 ,'ritm_type':row['cmdb_ci']
                                 ,'created_date':row['sys_created_on']
                                 ,'requested_for':row['requested_for']
                                 , 'short_description':row['short_description']
                                 , 'Project_Official_Title':''
                                 #, 'Primary_Project_Contact':row['Project']
                                 , 'PI_or_Primary_Contact':''
                                 , 'Project_is_currently_Active':'Yes'
                                 , 'irb':''
                                 })
            else:
                ##get info from word document
                p = merge_df['path'].to_string(index=False).lstrip()
                qd_df = process_word_doc(p)
                print(qd_df['Study Title'])
                ##write out
                out_info.append({'ritm':row['SN_RITM']
                                 ,'ritm_type':row['cmdb_ci']
                                 ,'created_date':row['sys_created_on']
                                 ,'requested_for':row['requested_for']
                                 , 'short_description':row['short_description']
                                 , 'Project_Official_Title': qd_df.get('Study Title', '') #or ''  #qd_df['Study Title']
                                 ##, 'Primary_Project_Contact':row['Project']
                                 , 'PI_or_Primary_Contact': qd_df.get('Principal Investigator', '') #or ''      #qd_df['Principal Investigator']
                                 , 'Project_is_currently_Active':'Yes' 
                                 , 'irb': qd_df.get('IRB Submission ID number', '') #or ''  #qd_df['IRB Submission ID number']
                                 })

        except Exception as ex:
            template = "An exception of type {0} occurred. Arguments:\n{1!r}"
            message = template.format(type(ex).__name__, ex.args)
            print(message)
            try:
                ## Try processing the document as the new format of Report Spec.
                rs = read_report_spec_v2.process_word_doc(p)
                out_info.append({'ritm':row['SN_RITM']
                                 ,'ritm_type':row['cmdb_ci']
                                 ,'created_date':row['sys_created_on']
                                 ,'requested_for':row['requested_for']
                                 , 'short_description':row['short_description']
                                 , 'Project_Official_Title':rs['Study Title'] #or ''  #qd_df['Study Title']
                                 , 'Primary_Project_Contact':rs['Point-of-contact']
                                 , 'PI_or_Primary_Contact': rs['Principal Investigator'] #or ''      #qd_df['Principal Investigator']
                                 , 'Project_is_currently_Active':'Yes' 
                                 , 'irb': rs['IRB Submission ID number'] #or ''  #qd_df['IRB Submission ID number']
                                 })
                print(rs['SN_RITM'])
                
            except Exception as ex:
                pass
                message = template.format(type(ex).__name__, ex.args)
                print(message)
            
            pass


    ##write out out_info
    
    print('Done!')
    out_info_df = pd.DataFrame(out_info)
    out_info_df.to_csv(qd_details, index=False)
    
    
    
if __name__ == "__main__":
    main()